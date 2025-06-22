import os
import pypandoc
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

import mysql.connector
from dotenv import load_dotenv

# Carga de credenciales de conexión desde el archivo .env externo
dotenv_path = os.path.join(os.path.dirname(__file__), ".env.bbdd")
load_dotenv(dotenv_path=dotenv_path)

MYSQL_CONFIG = {
    "host": os.getenv("MYSQL_HOST"),
    "user": os.getenv("MYSQL_USER"),
    "password": os.getenv("MYSQL_PASSWORD"),
    "database": os.getenv("MYSQL_DATABASE")
}


# ----------------------------- CONFIGURACIÓN -----------------------------
class ConfiguracionPresupuesto:
    def __init__(self):
        self.nombre_empresa = "EMPRESA DEMO"
        self.cif = "B12345678"
        self.telefono = "900 123 456"
        self.email = "info@empresademo.com"
        self.web = "www.empresademo.com"
        self.logo_path = "empresa_demo.png"
        self.titulo_documento = "PRESUPUESTO"
        self.color_principal = RGBColor(0, 83, 156)
        self.color_secundario = "0053C9"
        self.color_sombreado_encabezado = "D9D9D9"
        self.color_sombreado_total = "E6E6E6"
        self.porcentaje_iva = 21
        self.margen_superior = 2
        self.margen_inferior = 2
        self.margen_izquierdo = 2.5
        self.margen_derecho = 2.5
        self.titulo_datos_cliente = "DATOS DEL CLIENTE"
        self.titulo_descripcion = "DESCRIPCIÓN DEL TRABAJO"
        self.titulo_conceptos = "ESTIMACIÓN"
        self.titulo_condiciones = "CONDICIONES"
        self.titulo_aceptacion = "ACEPTACIÓN DEL PRESUPUESTO"
        self.encabezados_conceptos = ["CONCEPTO", "CANTIDAD", "PRECIO UNITARIO", "IMPORTE"]
        self.texto_subtotal = "Subtotal"
        self.texto_iva = f"IVA ({self.porcentaje_iva}%)"
        self.texto_total = "TOTAL"
        self.texto_aceptacion = "Mediante la firma del presente documento, el cliente acepta las condiciones y autoriza la realización de los trabajos descritos."
        self.texto_firma_empresa = "Firma de la empresa:"
        self.texto_firma_cliente = "Firma del cliente:"
        self.condiciones = []


# --------------------------- FUNCIONES ---------------------------

# Recupera de la base de datos todos los datos de un cliente dado su ID
def obtener_cliente_por_id(id_cliente):
    conn = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clientes WHERE id_cliente = %s", (id_cliente,))
    cliente = cursor.fetchone()
    cursor.close()
    conn.close()
    return cliente

# Genera un nuevo ID de trabajo, asegurando que no se repita con trabajos anteriores
def generar_nuevo_id_trabajo():
    conn = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor()
    cursor.execute("SELECT trabajos_asociados FROM clientes")
    trabajos = []
    for fila in cursor.fetchall():
        if fila[0]:
            trabajos += [int(x) for x in fila[0].split(",") if x.strip().isdigit()]
    nuevo_id = max(trabajos) + 1 if trabajos else 1
    cursor.close()
    conn.close()
    return nuevo_id

# Asocia un nuevo ID de trabajo a un cliente en la base de datos, actualizando su historial de trabajos
def asociar_trabajo_a_cliente(id_cliente, id_trabajo):
    conn = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor()
    cursor.execute("SELECT trabajos_asociados FROM clientes WHERE id_cliente = %s", (id_cliente,))
    fila = cursor.fetchone()
    trabajos = fila[0] or ""
    nuevos = trabajos.split(",") if trabajos else []
    nuevos.append(str(id_trabajo))
    nuevos_unicos = ",".join(sorted(set(nuevos), key=int))
    cursor.execute("UPDATE clientes SET trabajos_asociados = %s WHERE id_cliente = %s", (nuevos_unicos, id_cliente))
    conn.commit()
    cursor.close()
    conn.close()

# Aplica un color de fondo hexadecimal a una celda de tabla en un documento Word
def sombrear_celda(celda, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    celda._tc.get_or_add_tcPr().append(shading_elm)

# Aplica el estilo visual definido para títulos secundarios de secciones en el documento
def formato_titulo(parrafo):
    parrafo.style = "Heading 2"
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after = Pt(0)
    for run in parrafo.runs:
        run.font.color.rgb = RGBColor(0, 83, 156)
        run.font.size = Pt(12)
        run.font.name = "Calibri"

# Aplica el estilo visual principal al título del documento (centrado, grande y corporativo)
def formato_titulo_principal(parrafo):
    run = parrafo.runs[0]
    run.font.size = Pt(26)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 83, 156)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after = Pt(0)

# Aplica espaciado vertical limpio a un párrafo, sin modificar su contenido
def formato_parrafo(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

# Convierte un documento Word (.docx) a PDF usando la herramienta Pandoc
def convertir_docx_a_pdf(ruta_docx):
    carpeta_destino = os.path.dirname(ruta_docx)
    nombre_pdf = os.path.splitext(os.path.basename(ruta_docx))[0] + ".pdf"
    ruta_pdf = os.path.join(carpeta_destino, nombre_pdf)
    try:
        output = pypandoc.convert_file(ruta_docx, 'pdf', outputfile=ruta_pdf)
        return ruta_pdf
    except Exception as e:
        return None

# Genera un presupuesto en Word y PDF a partir de los datos del cliente y pedido, y devuelve las rutas generadas
def generar_presupuesto(pedido, config, id_trabajo):
    condiciones = pedido.get("condiciones", [])
    if isinstance(condiciones, str):
        condiciones = [line.strip() for line in condiciones.splitlines() if line.strip()]
    pedido["condiciones"] = condiciones
    
    cliente = obtener_cliente_por_id(pedido["id_cliente"])

    doc = Document()
    for seccion in doc.sections:
        seccion.top_margin = Cm(config.margen_superior)
        seccion.bottom_margin = Cm(config.margen_inferior)
        seccion.left_margin = Cm(config.margen_izquierdo)
        seccion.right_margin = Cm(config.margen_derecho)

    if os.path.exists(config.logo_path):
        doc.add_picture(config.logo_path, width=Cm(5))

    # Título principal
    titulo = doc.add_heading(config.titulo_documento, level=1)
    formato_titulo_principal(titulo)

    # Fecha
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    p_fecha = doc.add_paragraph(f"Fecha: {fecha_actual}")
    formato_parrafo(p_fecha)
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # DATOS DEL CLIENTE
    doc.add_paragraph()
    t_datos = doc.add_paragraph(config.titulo_datos_cliente)
    t_datos.runs[0].bold = True
    formato_titulo(t_datos)

    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = 'Table Grid'
    tabla.allow_autofit = False
    datos_cliente = [
        ("Nombre", cliente["nombre"]),
        ("DNI/CIF", cliente["dni"]),
        ("Teléfono", cliente["telefono"]),
        ("Correo", cliente["correo"] or "-"),
        ("Dirección", cliente["direccion"]),
        ("Población", f"{cliente['poblacion']}, {cliente['cp']} ({cliente['provincia']})")
    ]
    for i, (k, v) in enumerate(datos_cliente):
        tabla.cell(i, 0).text = k
        tabla.cell(i, 1).text = v

    # DESCRIPCIÓN DEL TRABAJO
    doc.add_paragraph()
    t_desc = doc.add_paragraph(config.titulo_descripcion)
    t_desc.runs[0].bold = True
    formato_titulo(t_desc)

    p_desc = doc.add_paragraph(pedido.get("descripcion", "-"))
    formato_parrafo(p_desc)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ESTIMACIÓN
    doc.add_paragraph()
    t_est = doc.add_paragraph(config.titulo_conceptos)
    t_est.runs[0].bold = True
    formato_titulo(t_est)

    tabla_c = doc.add_table(rows=1, cols=4)
    tabla_c.style = "Table Grid"
    for i, encabezado in enumerate(config.encabezados_conceptos):
        cell = tabla_c.cell(0, i)
        cell.text = encabezado
        cell.paragraphs[0].runs[0].bold = True
        sombrear_celda(cell, config.color_sombreado_encabezado)

    subtotal = 0
    for item in pedido["conceptos"]:
        fila = tabla_c.add_row().cells
        fila[0].text = item["concepto"]
        fila[1].text = str(item["cantidad"])
        fila[2].text = f"{item['precio_unitario']:.2f} €"
        importe = item["cantidad"] * item["precio_unitario"]
        fila[3].text = f"{importe:.2f} €"
        subtotal += importe

    iva = subtotal * (config.porcentaje_iva / 100)
    total = subtotal + iva

    for texto, valor, sombrear in [
        (config.texto_subtotal, f"{subtotal:.2f} €", False),
        (config.texto_iva, f"{iva:.2f} €", False),
        (config.texto_total, f"{total:.2f} €", True)
    ]:
        fila = tabla_c.add_row().cells
        fila[0].merge(fila[1]).merge(fila[2])
        fila[0].text = texto
        fila[3].text = valor
        if sombrear:
            sombrear_celda(fila[0], config.color_sombreado_total)
            sombrear_celda(fila[3], config.color_sombreado_total)

    # CONDICIONES
    doc.add_paragraph()
    t_cond = doc.add_paragraph(config.titulo_condiciones)
    t_cond.runs[0].bold = True
    formato_titulo(t_cond)

    for c in pedido.get("condiciones", config.condiciones):
        p = doc.add_paragraph(c, style='List Bullet')
        formato_parrafo(p)

    # ACEPTACIÓN
    doc.add_paragraph()
    t_acept = doc.add_paragraph(config.titulo_aceptacion)
    t_acept.runs[0].bold = True
    formato_titulo(t_acept)

    p_acept = doc.add_paragraph(config.texto_aceptacion)
    formato_parrafo(p_acept)

    # Firmas
    doc.add_paragraph()
    tabla_firmas = doc.add_table(rows=1, cols=2)
    tabla_firmas.style = 'Table Grid'
    firma_empresa = tabla_firmas.cell(0, 0)
    firma_empresa.text = config.texto_firma_empresa
    for _ in range(5): firma_empresa.add_paragraph()
    firma_cliente = tabla_firmas.cell(0, 1)
    firma_cliente.text = config.texto_firma_cliente
    for _ in range(5): firma_cliente.add_paragraph()

    # Footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{config.nombre_empresa} • CIF: {config.cif} • Tel: {config.telefono} • Email: {config.email} • {config.web}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    # Ruta docx
    base_dir = os.path.dirname(os.path.abspath(__file__))  # Ruta absoluta a Scripts/
    
    carpeta_destino = os.path.join(base_dir, "Data", "Presupuestos", "Pendiente")
    os.makedirs(carpeta_destino, exist_ok=True)

    ruta_docx_abs = pedido.get("ruta_guardado", os.path.join(carpeta_destino, f"Presupuesto_{id_trabajo}.docx"))
    doc.save(ruta_docx_abs)

    ruta_pdf_abs = convertir_docx_a_pdf(ruta_docx_abs)

    # Normaliza rutas para evitar None
    if ruta_pdf_abs is None:
        ruta_pdf_abs = ""

    # Convierte rutas absolutas a relativas respecto a Scripts/
    ruta_docx_rel = os.path.relpath(ruta_docx_abs, base_dir).replace("\\", "/")
    ruta_pdf_rel = os.path.relpath(ruta_pdf_abs, base_dir).replace("\\", "/") if ruta_pdf_abs else ""

    return ruta_docx_rel, ruta_pdf_rel, id_trabajo