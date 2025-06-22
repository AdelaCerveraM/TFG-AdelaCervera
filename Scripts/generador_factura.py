import os
import pypandoc
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

import mysql.connector
from dotenv import load_dotenv

# Carga de credenciales de conexión desde el archivo .env externo
dotenv_path = os.path.join(os.path.dirname(__file__), ".env.bbdd")
load_dotenv(dotenv_path=dotenv_path)

MYSQL_CONFIG = {
    "host": os.getenv("MYSQL_HOST"),
    "user": os.getenv("MYSQL_USER"),
    "password": os.getenv("MYSQL_PASSWORD"),
    "database": os.getenv("MYSQL_DATABASE")
}


# ----------------------------- CONFIGURACIÓN -----------------------------
class ConfiguracionFactura:
    def __init__(self):
        self.nombre_empresa = "EMPRESA DEMO"
        self.cif = "B12345678"
        self.telefono = "900 123 456"
        self.email = "info@empresademo.com"
        self.web = "www.empresademo.com"
        self.logo_path = "empresa_demo.png"
        self.titulo_documento = "FACTURA"
        self.color_principal = RGBColor(0, 83, 156)
        self.color_secundario = "0053C9"
        self.color_sombreado_encabezado = "D9D9D9"
        self.color_sombreado_total = "E6E6E6"
        self.porcentaje_iva = 21
        self.margen_superior = 2
        self.margen_inferior = 2
        self.margen_izquierdo = 2.5
        self.margen_derecho = 2.5
        self.encabezados_conceptos = ["CONCEPTO", "CANTIDAD", "PRECIO UNITARIO", "IMPORTE"]
        self.texto_subtotal = "Subtotal"
        self.texto_iva = f"IVA ({self.porcentaje_iva}%)"
        self.texto_total = "TOTAL"
        self.texto_pago = "Forma de pago: transferencia bancaria a la cuenta ES00 0000 0000 0000 0000"
        self.titulo_condiciones = "INFORMACIÓN ADICIONAL"


# --------------------------- FUNCIONES ---------------------------

# Consulta los datos de un cliente en la base de datos a partir de su ID
def obtener_cliente_por_id(id_cliente):
    conn = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clientes WHERE id_cliente = %s", (id_cliente,))
    cliente = cursor.fetchone()
    cursor.close()
    conn.close()
    return cliente

# Aplica un color de fondo a una celda de tabla en un documento Word (DOCX)
def sombrear_celda(celda, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    celda._tc.get_or_add_tcPr().append(shading_elm)

# Convierte un archivo Word (.docx) a PDF utilizando pypandoc
def convertir_docx_a_pdf(ruta_docx):
    carpeta_destino = os.path.dirname(ruta_docx)
    nombre_pdf = os.path.splitext(os.path.basename(ruta_docx))[0] + ".pdf"
    ruta_pdf = os.path.join(carpeta_destino, nombre_pdf)
    try:
        pypandoc.convert_file(ruta_docx, 'pdf', outputfile=ruta_pdf)
        return ruta_pdf
    except Exception as e:
        return None

# Genera una factura en formato Word y PDF a partir de los datos del cliente y el pedido, aplicando la configuración personalizada
def generar_factura(pedido, config, id_trabajo):
    cliente = obtener_cliente_por_id(pedido["id_cliente"])
    doc = Document()

    for seccion in doc.sections:
        seccion.top_margin = Cm(config.margen_superior)
        seccion.bottom_margin = Cm(config.margen_inferior)
        seccion.left_margin = Cm(config.margen_izquierdo)
        seccion.right_margin = Cm(config.margen_derecho)

    if os.path.exists(config.logo_path):
        doc.add_picture(config.logo_path, width=Cm(5))

    doc.add_heading(config.titulo_documento, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cliente
    doc.add_paragraph("DATOS DEL CLIENTE", style='Heading 2')
    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = 'Table Grid'
    tabla.allow_autofit = False
    datos = [
        ("Nombre", cliente["nombre"]),
        ("DNI/CIF", cliente["dni"]),
        ("Teléfono", cliente["telefono"]),
        ("Correo", cliente["correo"]),
        ("Dirección", cliente["direccion"]),
        ("Población", f"{cliente['poblacion']}, {cliente['cp']} ({cliente['provincia']})")
    ]
    for i, (k, v) in enumerate(datos):
        tabla.cell(i, 0).text = str(k)
        tabla.cell(i, 1).text = str(v) if v is not None else ""

    # Conceptos
    doc.add_paragraph()
    doc.add_paragraph("DETALLE", style='Heading 2')
    tabla_c = doc.add_table(rows=1, cols=4)
    tabla_c.style = "Table Grid"
    for i, encabezado in enumerate(config.encabezados_conceptos):
        cell = tabla_c.cell(0, i)
        cell.text = encabezado
        sombrear_celda(cell, config.color_sombreado_encabezado)

    subtotal = 0
    for item in pedido["conceptos"]:
        fila = tabla_c.add_row().cells
        fila[0].text = item["concepto"]
        fila[1].text = str(item["cantidad"])
        fila[2].text = f"{item['precio_unitario']:.2f} €"
        importe = item["cantidad"] * item["precio_unitario"]
        fila[3].text = f"{importe:.2f} €"
        subtotal += importe

    iva = subtotal * (config.porcentaje_iva / 100)
    total = subtotal + iva

    for texto, valor, sombrear in [
        (config.texto_subtotal, f"{subtotal:.2f} €", False),
        (config.texto_iva, f"{iva:.2f} €", False),
        (config.texto_total, f"{total:.2f} €", True)
    ]:
        fila = tabla_c.add_row().cells
        fila[0].merge(fila[1]).merge(fila[2])
        fila[0].text = texto
        fila[3].text = valor
        if sombrear:
            sombrear_celda(fila[0], config.color_sombreado_total)
            sombrear_celda(fila[3], config.color_sombreado_total)

    # Información adicional
    doc.add_paragraph()
    doc.add_paragraph(config.titulo_condiciones, style='Heading 2')
    texto_info = pedido.get("informacion_adicional", config.texto_pago)
    texto_info = texto_info.replace('\\"', '"').strip('"')
    for linea in texto_info.splitlines():
        if linea.strip():
            doc.add_paragraph(linea.strip(), style='List Bullet')
    doc.add_paragraph(texto_info)

    # Footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{config.nombre_empresa} • CIF: {config.cif} • Tel: {config.telefono} • Email: {config.email} • {config.web}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    # Guardar
    base_dir = os.path.dirname(os.path.abspath(__file__))
    carpeta = os.path.join(base_dir, "Data", "Facturas", "Pendiente")
    os.makedirs(carpeta, exist_ok=True)

    ruta_docx = os.path.join(carpeta, f"Factura_{id_trabajo}.docx")
    doc.save(ruta_docx)
    ruta_pdf = convertir_docx_a_pdf(ruta_docx)

    ruta_docx_rel = os.path.relpath(ruta_docx, base_dir).replace("\\", "/")
    ruta_pdf_rel = os.path.relpath(ruta_pdf, base_dir).replace("\\", "/") if ruta_pdf else ""

    return ruta_docx_rel, ruta_pdf_rel, id_trabajo